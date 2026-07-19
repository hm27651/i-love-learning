from __future__ import annotations

import csv
import hashlib
import json
import re
import shutil
import sqlite3
import threading
from concurrent.futures import ThreadPoolExecutor
from datetime import datetime
from difflib import SequenceMatcher
from pathlib import Path, PurePath

from app_runtime import connect_database
from services.imports.pdf_question_parser import parse_pdf as parse_h3c_pdf
from transfer_service import PACKAGE_VERSION, inspect_share_package


PARSER_VERSION = "study-import-1.0"
SUPPORTED_TYPES = {"pdf", "docx", "xlsx", "csv", "zip"}
TYPE_ALIASES = {
    "单选": "single", "单选题": "single", "single": "single",
    "多选": "multiple", "多选题": "multiple", "multiple": "multiple",
    "判断": "true_false", "判断题": "true_false", "true_false": "true_false",
    "填空": "fill", "填空题": "fill", "fill": "fill",
    "简答": "short", "简答题": "short", "short": "short",
}
QUESTION_HEADER = re.compile(r"(?mi)^(?:QUESTION|问题|题目)?\s*(\d+)\s*[.、)]?\s*(.*)$")
ANSWER_LINE = re.compile(r"(?mi)^(?:Correct Answer|正确答案|参考答案|答案)\s*[:：]\s*(.+)$")
EXPLANATION_LINE = re.compile(r"(?mi)^(?:Explanation(?:/Reference)?|说明/参考|解析)\s*[:：]?\s*(.*)$")
OPTION_LINE = re.compile(r"(?m)^([A-H])[.、．:]\s*(.+)$")


def now_iso() -> str:
    return datetime.now().replace(microsecond=0).isoformat(sep=" ")


def normalize_text(value: str) -> str:
    return re.sub(r"\s+", "", (value or "")).casefold()


def fingerprint(stem: str, options: list[str]) -> str:
    raw = normalize_text(stem) + "\n" + "\n".join(normalize_text(x) for x in options)
    return hashlib.sha256(raw.encode("utf-8")).hexdigest()


def _answers(value: str, qtype: str) -> list[str]:
    value = (value or "").strip()
    if qtype in {"single", "multiple"}:
        found = re.findall(r"[A-H]", value.upper())
        return sorted(set(found)) if qtype == "multiple" else found[:1]
    if qtype == "true_false":
        return ["true" if value.casefold() in {"true", "1", "对", "正确", "是", "√"} else "false"]
    return [value] if value else []


def _infer_type(options: list[str], answer: str, explicit: str = "") -> str:
    if explicit.strip().casefold() in TYPE_ALIASES:
        return TYPE_ALIASES[explicit.strip().casefold()]
    letters = re.findall(r"[A-H]", (answer or "").upper())
    if options:
        return "multiple" if len(set(letters)) > 1 else "single"
    if (answer or "").strip().casefold() in {"true", "false", "对", "错", "正确", "错误", "是", "否", "√", "×"}:
        return "true_false"
    return "fill"


def _validate(record: dict) -> str:
    if not record.get("stem", "").strip():
        return "缺少题干"
    if not record.get("answer") or not any(str(x).strip() for x in record["answer"]):
        return "缺少答案"
    if record["type"] in {"single", "multiple"}:
        if len(record.get("options", [])) < 2:
            return "客观题缺少选项"
        labels = set("ABCDEFGH"[:len(record["options"])])
        if not set(record["answer"]).issubset(labels):
            return "答案与选项不一致"
    return ""


def parse_structured_text(text: str) -> tuple[list[dict], list[dict]]:
    lines = text.replace("\r\n", "\n").replace("\r", "\n")
    headers = list(re.finditer(r"(?mi)^(?:QUESTION|问题|题目)?\s*(\d+)\s*[.、)]\s*(.*)$", lines))
    if not headers:
        headers = list(re.finditer(r"(?mi)^(?:QUESTION|问题)\s+(\d+)\s*$", lines))
    records, anomalies = [], []
    for index, header in enumerate(headers):
        end = headers[index + 1].start() if index + 1 < len(headers) else len(lines)
        block = (header.group(2) + "\n" + lines[header.end():end]).strip()
        answer_match = ANSWER_LINE.search(block)
        if not answer_match:
            anomalies.append({"item": header.group(1), "reason": "缺少答案"})
            continue
        before = block[:answer_match.start()].strip()
        option_matches = list(OPTION_LINE.finditer(before))
        options = [match.group(2).strip() for match in option_matches]
        stem = before[:option_matches[0].start()].strip() if option_matches else before
        raw_answer = answer_match.group(1).strip()
        qtype = _infer_type(options, raw_answer)
        explanation_match = EXPLANATION_LINE.search(block, answer_match.end())
        explanation = explanation_match.group(1).strip() if explanation_match else ""
        record = {
            "source_page": None,
            "source_item_key": header.group(1),
            "type": qtype,
            "stem": stem,
            "options": options,
            "answer": _answers(raw_answer, qtype),
            "explanation": explanation,
            "chapter_name": "",
            "point_name": "",
        }
        record["validation_error"] = _validate(record)
        records.append(record)
    if not records and not anomalies:
        anomalies.append({"reason": "未识别到结构化题目；请使用固定模板或可复制的题号/选项/答案格式"})
    return records, anomalies


def parse_pdf(path: Path) -> tuple[list[dict], list[dict]]:
    records, anomalies, _ = parse_h3c_pdf(path)
    result = []
    for item in records:
        record = {
            "source_page": item.get("page"), "source_item_key": item.get("key", str(item.get("number", ""))),
            "type": item["type"], "stem": item["stem"], "options": item.get("options", []),
            "answer": item.get("answer", []), "explanation": item.get("explanation", ""),
            "chapter_name": "", "point_name": "",
        }
        record["validation_error"] = _validate(record)
        result.append(record)
    return result, anomalies


def parse_docx(path: Path) -> tuple[list[dict], list[dict]]:
    from docx import Document

    doc = Document(path)
    chunks = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            chunks.append("\t".join(cell.text for cell in row.cells))
    return parse_structured_text("\n".join(chunks))


def _record_from_mapping(row: dict[str, object], item_index: int) -> dict:
    def value(*names: str) -> str:
        for name in names:
            if name in row and row[name] is not None:
                return str(row[name]).strip()
        return ""
    options = [value(f"选项{letter}", f"option_{letter.lower()}") for letter in "ABCDEFGH"]
    while options and not options[-1]:
        options.pop()
    raw_answer = value("答案", "answer")
    qtype = _infer_type(options, raw_answer, value("题型", "type"))
    record = {
        "source_page": None, "source_item_key": value("原题号", "题号", "source_item_key") or str(item_index),
        "type": qtype, "stem": value("题干", "stem"), "options": options,
        "answer": _answers(raw_answer, qtype), "explanation": value("解析", "explanation"),
        "chapter_name": value("章节", "chapter"), "point_name": value("知识点", "knowledge_point"),
    }
    record["validation_error"] = _validate(record)
    return record


def parse_csv(path: Path) -> tuple[list[dict], list[dict]]:
    raw = path.read_bytes()
    text = None
    for encoding in ("utf-8-sig", "gb18030"):
        try:
            text = raw.decode(encoding)
            break
        except UnicodeDecodeError:
            pass
    if text is None:
        return [], [{"reason": "CSV 编码无法识别，请保存为 UTF-8"}]
    rows = list(csv.DictReader(text.splitlines()))
    return [_record_from_mapping(row, index) for index, row in enumerate(rows, 1)], []


def parse_xlsx(path: Path) -> tuple[list[dict], list[dict]]:
    from openpyxl import load_workbook

    workbook = load_workbook(path, read_only=True, data_only=True)
    sheet = workbook.active
    iterator = sheet.iter_rows(values_only=True)
    headers = [str(value).strip() if value is not None else "" for value in next(iterator, [])]
    rows = [_record_from_mapping(dict(zip(headers, values)), index) for index, values in enumerate(iterator, 1)]
    workbook.close()
    return rows, []


def parse_document(path: Path, file_type: str) -> tuple[list[dict], list[dict]]:
    if file_type == "pdf":
        return parse_pdf(path)
    if file_type == "docx":
        return parse_docx(path)
    if file_type == "xlsx":
        return parse_xlsx(path)
    if file_type == "csv":
        return parse_csv(path)
    raise ValueError("不支持的文件格式")


def probe_document(path: Path, file_type: str) -> str:
    """Read only a small first-page/header sample for target suggestions."""
    if file_type == "pdf":
        import pdfplumber
        with pdfplumber.open(path) as pdf:
            return (pdf.pages[0].extract_text() or "")[:4000] if pdf.pages else ""
    if file_type == "docx":
        from docx import Document
        document = Document(path)
        return "\n".join(p.text for p in document.paragraphs[:30])[:4000]
    if file_type == "xlsx":
        from openpyxl import load_workbook
        workbook = load_workbook(path, read_only=True, data_only=True)
        rows = []
        for values in workbook.active.iter_rows(max_row=5, values_only=True):
            rows.append(" ".join(str(value) for value in values if value is not None))
        workbook.close()
        return "\n".join(rows)[:4000]
    if file_type == "csv":
        return path.read_bytes()[:8000].decode("utf-8-sig", errors="ignore")[:4000]
    return ""


def _connect(db_path: Path) -> sqlite3.Connection:
    return connect_database(db_path)


def detect_job(db_path: Path, data_dir: Path, job_id: str) -> None:
    conn = _connect(db_path)
    try:
        conn.execute("UPDATE import_jobs SET status='running',stage='detecting',progress=10,updated_at=? WHERE id=?", (now_iso(), job_id))
        conn.commit()
        row = conn.execute("""SELECT j.*,d.original_name,d.stored_path,d.file_type FROM import_jobs j
          JOIN source_documents d ON d.id=j.source_document_id WHERE j.id=?""", (job_id,)).fetchone()
        if row["file_type"] == "zip":
            manifest, questions = inspect_share_package(data_dir / row["stored_path"])
            project_name = manifest.get("project", {}).get("name", "")
            project_ids = [item["id"] for item in conn.execute(
                "SELECT id,name FROM learning_projects WHERE status='active'"
            ) if normalize_text(item["name"]) == normalize_text(project_name)]
            detected = {
                "project_ids": project_ids,
                "reason": "已识别我爱学习题库分享包，请选择目标项目",
                "package": {
                    "version": manifest["version"], "project": manifest["project"], "scope": manifest["scope"],
                    "counts": manifest["counts"], "subjects": manifest["subjects"],
                    "warnings": manifest.get("warnings", []),
                },
            }
            conn.execute("""UPDATE import_jobs SET import_kind='package',status='waiting_target',stage='target',
              progress=20,detected_json=?,package_json=?,message='分享包已验证，请选择目标项目',updated_at=? WHERE id=?""",
              (json.dumps(detected, ensure_ascii=False), json.dumps(manifest, ensure_ascii=False), now_iso(), job_id))
            conn.commit()
            return
        filename = row["original_name"]
        probe = probe_document(data_dir / row["stored_path"], row["file_type"])
        haystack = filename + "\n" + probe
        suggestions = {"project_ids": [], "subject_ids": [], "reason": "根据文件名和首页内容建议，请人工确认", "probe_excerpt": probe[:300]}
        for project in conn.execute("SELECT id,name FROM learning_projects WHERE status='active'"):
            if normalize_text(project["name"]) in normalize_text(haystack): suggestions["project_ids"].append(project["id"])
        for subject in conn.execute("SELECT id,project_id,name,code FROM subjects"):
            if (subject["code"] and normalize_text(subject["code"]) in normalize_text(haystack)) or normalize_text(subject["name"]) in normalize_text(haystack):
                suggestions["subject_ids"].append(subject["id"])
                if subject["project_id"] not in suggestions["project_ids"]: suggestions["project_ids"].append(subject["project_id"])
        conn.execute("""UPDATE import_jobs SET status='waiting_target',stage='target',progress=20,
          detected_json=?,message='文件已保存，请确认项目与科目',updated_at=? WHERE id=?""",
          (json.dumps(suggestions, ensure_ascii=False), now_iso(), job_id))
        conn.commit()
    except Exception as exc:
        conn.rollback()
        conn.execute("UPDATE import_jobs SET status='failed',stage='detecting',message=?,error_json=?,updated_at=? WHERE id=?",
                     (str(exc), json.dumps([{"reason": str(exc)}], ensure_ascii=False), now_iso(), job_id))
        conn.commit()
    finally:
        conn.close()


def parse_job(db_path: Path, data_dir: Path, job_id: str) -> None:
    conn = _connect(db_path)
    try:
        conn.execute("DELETE FROM import_candidates WHERE job_id=?", (job_id,))
        conn.execute("UPDATE import_jobs SET status='running',stage='parsing',progress=30,message='正在解析题目',updated_at=? WHERE id=?", (now_iso(), job_id))
        conn.commit()
        job = conn.execute("""SELECT j.*,d.stored_path,d.file_type FROM import_jobs j JOIN source_documents d
          ON d.id=j.source_document_id WHERE j.id=?""", (job_id,)).fetchone()
        if job["file_type"] == "zip":
            _parse_share_package(conn, data_dir, job)
            return
        records, anomalies = parse_document(data_dir / job["stored_path"], job["file_type"])
        existing = []
        for row in conn.execute("""SELECT q.id,q.stem,q.options_json FROM questions q JOIN knowledge_points kp ON kp.id=q.knowledge_point_id
          JOIN chapters c ON c.id=kp.chapter_id WHERE c.subject_id=?""", (job["subject_id"],)):
            options = json.loads(row["options_json"])
            existing.append((row["id"], normalize_text(row["stem"]), fingerprint(row["stem"], options)))
        duplicate_count = valid_count = 0
        preview = []
        for index, record in enumerate(records, 1):
            fp = fingerprint(record["stem"], record.get("options", []))
            duplicate_id = next((qid for qid, _, old_fp in existing if old_fp == fp), None)
            if duplicate_id is None:
                normalized = normalize_text(record["stem"])
                duplicate_id = next((qid for qid, old_stem, _ in existing if len(normalized) >= 20 and SequenceMatcher(None, normalized, old_stem).ratio() >= .94), None)
            error = record.get("validation_error", "")
            valid_count += int(not error)
            duplicate_count += int(duplicate_id is not None)
            decision = "skip" if duplicate_id is not None else "insert"
            conn.execute("""INSERT INTO import_candidates(job_id,item_index,source_page,source_item_key,question_type,
              stem,options_json,answer_json,explanation,chapter_name,point_name,fingerprint,duplicate_question_id,
              validation_error,decision) VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
              (job_id, index, record.get("source_page"), record.get("source_item_key", ""), record["type"], record["stem"],
               json.dumps(record.get("options", []), ensure_ascii=False), json.dumps(record.get("answer", []), ensure_ascii=False),
               record.get("explanation", ""), record.get("chapter_name", ""), record.get("point_name", ""), fp,
               duplicate_id, error, decision))
            if len(preview) < 20:
                preview.append({"index": index, "type": record["type"], "stem": record["stem"][:160], "error": error, "duplicate": duplicate_id})
        blocking_reasons = {"blank_answer", "invalid_answer", "blank_stem", "缺少答案", "缺少题干", "答案与选项不一致"}
        blocking_anomalies = [item for item in anomalies if item.get("reason") in blocking_reasons or "未识别" in str(item.get("reason", ""))]
        status = "blocked" if any(record.get("validation_error") for record in records) or blocking_anomalies or not records else "ready"
        message = "存在必填错误，整批暂不可提交" if status == "blocked" else "解析完成，请处理重复项后提交"
        conn.execute("""UPDATE import_jobs SET status=?,stage='review',progress=90,message=?,parser_version=?,
          error_json=?,preview_json=?,candidate_count=?,valid_count=?,duplicate_count=?,updated_at=? WHERE id=?""",
          (status, message, PARSER_VERSION, json.dumps(anomalies, ensure_ascii=False), json.dumps(preview, ensure_ascii=False),
           len(records), valid_count, duplicate_count, now_iso(), job_id))
        conn.commit()
    except Exception as exc:
        conn.rollback()
        conn.execute("UPDATE import_jobs SET status='failed',stage='parsing',message=?,error_json=?,updated_at=? WHERE id=?",
                     (str(exc), json.dumps([{"reason": str(exc)}], ensure_ascii=False), now_iso(), job_id))
        conn.commit()
    finally:
        conn.close()


def _parse_share_package(conn: sqlite3.Connection, data_dir: Path, job) -> None:
    mapping = json.loads(job["mapping_json"] or "{}")
    subject_mapping = mapping.get("subjects", {})
    if not job["project_id"] or not subject_mapping:
        raise ValueError("请先确认项目和科目映射")
    staging = data_dir / "imports" / "staging" / job["id"]
    if staging.exists():
        shutil.rmtree(staging)
    manifest, records = inspect_share_package(data_dir / job["stored_path"], staging)
    subject_meta = {item["key"]: item for item in manifest["subjects"]}
    chapter_meta = {item["key"]: item for subject in manifest["subjects"] for item in subject["chapters"]}
    point_meta = {item["key"]: item for subject in manifest["subjects"] for chapter in subject["chapters"] for item in chapter["points"]}

    existing_cache: dict[int, list[tuple[int, str, str]]] = {}
    for value in subject_mapping.values():
        subject_id = value.get("target_subject_id") if isinstance(value, dict) else None
        if not subject_id or subject_id in existing_cache:
            continue
        rows = []
        for row in conn.execute("""SELECT q.id,q.stem,q.options_json FROM questions q
          JOIN knowledge_points kp ON kp.id=q.knowledge_point_id JOIN chapters c ON c.id=kp.chapter_id
          WHERE c.subject_id=?""", (subject_id,)):
            options = json.loads(row["options_json"])
            rows.append((row["id"], normalize_text(row["stem"]), fingerprint(row["stem"], options)))
        existing_cache[subject_id] = rows

    duplicate_count = valid_count = 0
    preview = []
    for index, record in enumerate(records, 1):
        subject = subject_meta[record["subject_key"]]
        chapter = chapter_meta[record["chapter_key"]]
        point = point_meta[record["point_key"]]
        mapped = subject_mapping.get(record["subject_key"], {})
        subject_id = mapped.get("target_subject_id") if isinstance(mapped, dict) else None
        fp = fingerprint(record["stem"], record.get("options", []))
        duplicate_id = None
        existing = existing_cache.get(subject_id, [])
        if existing:
            duplicate_id = next((qid for qid, _, old_fp in existing if old_fp == fp), None)
            if duplicate_id is None:
                normalized = normalize_text(record["stem"])
                duplicate_id = next((qid for qid, old_stem, _ in existing
                                     if len(normalized) >= 20 and SequenceMatcher(None, normalized, old_stem).ratio() >= .94), None)
        error = _validate(record)
        valid_count += int(not error)
        duplicate_count += int(duplicate_id is not None)
        conn.execute("""INSERT INTO import_candidates(job_id,item_index,source_item_key,question_type,stem,
          options_json,answer_json,explanation,chapter_name,point_name,fingerprint,duplicate_question_id,
          validation_error,decision,subject_key,subject_name,subject_code,chapter_key,chapter_is_core,
          point_key,difficulty,sender_status,image_ref,image_missing)
          VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
          (job["id"], index, record["item_key"], record["type"], record["stem"],
           json.dumps(record["options"], ensure_ascii=False), json.dumps(record["answer"], ensure_ascii=False),
           record["explanation"], chapter["name"], point["name"], fp, duplicate_id, error,
           "skip" if duplicate_id else "insert", subject["key"], subject["name"], subject.get("code", ""),
           chapter["key"], int(bool(chapter.get("is_core"))), point["key"], record["difficulty"],
           record["status"], record.get("image_ref", ""), int(bool(record.get("image_missing")))))
        if len(preview) < 20:
            preview.append({"index": index, "type": record["type"], "stem": record["stem"][:160],
                            "error": error, "duplicate": duplicate_id, "image_missing": record.get("image_missing", False)})
    status = "blocked" if any(item["validation_error"] for item in conn.execute(
        "SELECT validation_error FROM import_candidates WHERE job_id=?", (job["id"],)
    )) or not records else "ready"
    warnings = list(manifest.get("warnings", []))
    message = "存在必填错误，整批暂不可提交" if status == "blocked" else "分享包解析完成，请处理重复项后提交"
    conn.execute("""UPDATE import_jobs SET status=?,stage='review',progress=90,message=?,parser_version=?,
      error_json=?,preview_json=?,candidate_count=?,valid_count=?,duplicate_count=?,updated_at=? WHERE id=?""",
      (status, message, f"share-package-{PACKAGE_VERSION}", json.dumps(warnings, ensure_ascii=False),
       json.dumps(preview, ensure_ascii=False), len(records), valid_count, duplicate_count, now_iso(), job["id"]))
    conn.commit()


def ensure_target_point(
    conn: sqlite3.Connection,
    subject_id: int,
    chapter_name: str = "",
    point_name: str = "",
    is_core: bool = False,
) -> int:
    chapter_name = chapter_name.strip() or "待分类"
    point_name = point_name.strip() or "待分类"
    chapter = next((row for row in conn.execute("SELECT id,name,is_core FROM chapters WHERE subject_id=?", (subject_id,))
                    if normalize_text(row["name"]) == normalize_text(chapter_name)), None)
    chapter_id = chapter["id"] if chapter else conn.execute(
        "INSERT INTO chapters(subject_id,name,is_core) VALUES (?,?,?)", (subject_id, chapter_name, int(is_core))
    ).lastrowid
    if chapter and is_core and not chapter["is_core"]:
        conn.execute("UPDATE chapters SET is_core=1 WHERE id=?", (chapter_id,))
    point = next((row for row in conn.execute("SELECT id,name FROM knowledge_points WHERE chapter_id=?", (chapter_id,))
                  if normalize_text(row["name"]) == normalize_text(point_name)), None)
    return int(point["id"] if point else conn.execute(
        "INSERT INTO knowledge_points(chapter_id,name) VALUES (?,?)", (chapter_id, point_name)
    ).lastrowid)


def _commit_share_package(
    conn: sqlite3.Connection,
    job,
    candidates,
    decisions: dict[int, str],
    duplicate_action: str,
    status_strategy: str,
    data_dir: Path,
) -> dict:
    mapping = json.loads(job["mapping_json"] or "{}").get("subjects", {})
    if not mapping:
        raise ValueError("分享包尚未完成科目映射")
    subject_meta: dict[str, tuple[str, str]] = {}
    for candidate in candidates:
        subject_meta.setdefault(candidate["subject_key"], (candidate["subject_name"], candidate["subject_code"]))
    subject_ids: dict[str, int] = {}
    used_targets: set[int] = set()
    for key, (name, code) in subject_meta.items():
        item = mapping.get(key, {})
        target_id = item.get("target_subject_id") if isinstance(item, dict) else None
        if target_id:
            row = conn.execute("SELECT id FROM subjects WHERE id=? AND project_id=?", (target_id, job["project_id"])).fetchone()
            if not row:
                raise ValueError(f"科目“{name}”的映射目标无效")
            if target_id in used_targets:
                raise ValueError("不同来源科目不能映射到同一个目标科目")
            subject_ids[key] = int(target_id)
            used_targets.add(int(target_id))
        else:
            if code and conn.execute("SELECT 1 FROM subjects WHERE project_id=? AND code=?", (job["project_id"], code)).fetchone():
                raise ValueError(f"科目代码“{code}”已存在，请重新选择映射")
            subject_ids[key] = int(conn.execute(
                "INSERT INTO subjects(project_id,name,code) VALUES (?,?,?)", (job["project_id"], name, code)
            ).lastrowid)

    inserted = skipped = updated = 0
    now = now_iso()
    staging = data_dir / "imports" / "staging" / job["id"]
    uploads = data_dir / "uploads"
    uploads.mkdir(parents=True, exist_ok=True)
    preserve = status_strategy == "preserve"
    created_images: list[Path] = []
    try:
        for candidate in candidates:
            decision = decisions.get(candidate["id"], duplicate_action if candidate["duplicate_question_id"] else "insert")
            if decision not in {"skip", "update", "copy", "insert"}:
                decision = "skip" if candidate["duplicate_question_id"] else "insert"
            if candidate["duplicate_question_id"] and decision == "skip":
                skipped += 1
                continue
            subject_id = subject_ids[candidate["subject_key"]]
            point_id = ensure_target_point(conn, subject_id, candidate["chapter_name"], candidate["point_name"], bool(candidate["chapter_is_core"]))
            image_path = None
            if candidate["image_ref"] and not candidate["image_missing"]:
                staged = staging / PurePath(candidate["image_ref"])
                if not staged.is_file():
                    raise ValueError(f"题目 #{candidate['item_index']} 的图片暂存文件缺失")
                digest = Path(candidate["image_ref"]).stem
                image_path = f"share_{digest}{staged.suffix.lower()}"
                final_image = uploads / image_path
                if not final_image.exists():
                    shutil.copy2(staged, final_image)
                    created_images.append(final_image)
            status = candidate["sender_status"] if preserve and candidate["sender_status"] in {"verified", "draft"} else "draft"
            common_values = (
                point_id, candidate["question_type"], candidate["stem"], candidate["options_json"],
                candidate["answer_json"], candidate["explanation"], candidate["difficulty"], "", "",
                image_path, status, job["source_document_id"], None, "", job["id"],
                job["parser_version"] or f"share-package-{PACKAGE_VERSION}",
            )
            if candidate["duplicate_question_id"] and decision == "update":
                conn.execute("""UPDATE questions SET knowledge_point_id=?,type=?,stem=?,options_json=?,answer_json=?,
                  explanation=?,difficulty=?,source=?,version_note=?,image_path=COALESCE(?,image_path),status=?,source_document_id=?,
                  source_page=?,source_item_key=?,import_batch_id=?,parser_version=?,updated_at=? WHERE id=?""",
                  common_values + (now, candidate["duplicate_question_id"]))
                updated += 1
            else:
                conn.execute("""INSERT INTO questions(knowledge_point_id,type,stem,options_json,answer_json,explanation,
                  difficulty,source,version_note,image_path,status,source_document_id,source_page,source_item_key,
                  import_batch_id,parser_version,created_at,updated_at) VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
                  common_values + (now, now))
                inserted += 1
    except Exception:
        for image in created_images:
            image.unlink(missing_ok=True)
        raise
    written = inserted + updated
    mode = "保留发送方状态" if preserve else "全部转为草稿"
    message = f"已写入 {written} 道题，跳过 {skipped} 道重复题；{mode}"
    conn.execute("""UPDATE import_jobs SET status='committed',stage='complete',progress=100,message=?,
      committed_count=?,completed_at=?,updated_at=? WHERE id=?""", (message, written, now, now, job["id"]))
    try:
        conn.commit()
    except Exception:
        for image in created_images:
            image.unlink(missing_ok=True)
        raise
    if staging.exists():
        shutil.rmtree(staging, ignore_errors=True)
    return {"inserted": inserted, "updated": updated, "skipped": skipped, "written": written, "message": message}


def commit_job(
    conn: sqlite3.Connection,
    job_id: str,
    decisions: dict[int, str] | None = None,
    duplicate_action: str = "skip",
    status_strategy: str = "draft",
    data_dir: Path | None = None,
) -> dict:
    """Atomically commit one validated import batch. Used by both web and CLI entry points."""
    decisions = decisions or {}
    if duplicate_action not in {"skip", "update", "copy"}:
        duplicate_action = "skip"
    job = conn.execute("""SELECT j.*,d.original_name FROM import_jobs j JOIN source_documents d
      ON d.id=j.source_document_id WHERE j.id=?""", (job_id,)).fetchone()
    if not job or job["status"] != "ready" or (job["import_kind"] != "package" and not job["subject_id"]):
        raise ValueError("该导入任务尚未准备好")
    invalid = conn.execute(
        "SELECT COUNT(*) n FROM import_candidates WHERE job_id=? AND validation_error<>''", (job_id,)
    ).fetchone()["n"]
    if invalid:
        raise ValueError(f"还有 {invalid} 道题存在必填错误，整批未写入")
    candidates = conn.execute("SELECT * FROM import_candidates WHERE job_id=? ORDER BY item_index", (job_id,)).fetchall()
    if not candidates:
        raise ValueError("导入任务没有可提交题目")
    if job["import_kind"] == "package":
        if data_dir is None:
            raise ValueError("分享包导入缺少数据目录")
        try:
            return _commit_share_package(conn, job, candidates, decisions, duplicate_action, status_strategy, data_dir)
        except Exception:
            conn.rollback()
            uploads = data_dir / "uploads"
            for candidate in candidates:
                if not candidate["image_ref"]:
                    continue
                ref = PurePath(candidate["image_ref"])
                filename = f"share_{ref.stem}{ref.suffix.lower()}"
                if not conn.execute("SELECT 1 FROM questions WHERE image_path=? LIMIT 1", (filename,)).fetchone():
                    (uploads / filename).unlink(missing_ok=True)
            raise
    inserted = skipped = updated = 0
    now = now_iso()
    for candidate in candidates:
        decision = decisions.get(candidate["id"], duplicate_action if candidate["duplicate_question_id"] else "insert")
        if candidate["duplicate_question_id"] and decision == "skip":
            skipped += 1
            continue
        point_id = ensure_target_point(conn, job["subject_id"], candidate["chapter_name"], candidate["point_name"])
        source = f"{job['original_name']}；原题号 {candidate['source_item_key']}"
        if candidate["source_page"]:
            source += f"；第 {candidate['source_page']} 页"
        values = (
            point_id, candidate["question_type"], candidate["stem"], candidate["options_json"],
            candidate["answer_json"], candidate["explanation"], source,
            "通用导入中心生成的草稿，核对答案和版式后再标记为已核验。", job["source_document_id"],
            candidate["source_page"], candidate["source_item_key"], job_id,
            job["parser_version"] or PARSER_VERSION, now,
        )
        if candidate["duplicate_question_id"] and decision == "update":
            conn.execute("""UPDATE questions SET knowledge_point_id=?,type=?,stem=?,options_json=?,answer_json=?,
              explanation=?,source=?,version_note=?,source_document_id=?,source_page=?,source_item_key=?,
              import_batch_id=?,parser_version=?,status='draft',updated_at=? WHERE id=?""",
              values + (candidate["duplicate_question_id"],))
            updated += 1
        else:
            conn.execute("""INSERT INTO questions(knowledge_point_id,type,stem,options_json,answer_json,explanation,
              difficulty,source,version_note,status,source_document_id,source_page,source_item_key,import_batch_id,
              parser_version,created_at,updated_at) VALUES (?,?,?,?,?,?,2,?,?,'draft',?,?,?,?,?,?,?)""",
              values[:-1] + (now, now))
            inserted += 1
    written = inserted + updated
    message = f"已写入 {written} 道草稿，跳过 {skipped} 道重复题"
    conn.execute("""UPDATE import_jobs SET status='committed',stage='complete',progress=100,message=?,
      committed_count=?,completed_at=?,updated_at=? WHERE id=?""", (message, written, now, now, job_id))
    return {"inserted": inserted, "updated": updated, "skipped": skipped, "written": written, "message": message}


class ImportQueue:
    def __init__(self, db_path: Path, data_dir: Path, executor=None):
        self.db_path = db_path
        self.data_dir = data_dir
        self.executor = executor or ThreadPoolExecutor(max_workers=1, thread_name_prefix="study-import")
        self.lock = threading.Lock()

    def recover(self) -> None:
        conn = _connect(self.db_path)
        try:
            conn.execute("""UPDATE import_jobs SET status='interrupted',message='服务重启导致任务中断，可从原文件重试',
              updated_at=? WHERE status IN ('queued','running')""", (now_iso(),))
            conn.commit()
        finally:
            conn.close()

    def submit_detect(self, job_id: str) -> None:
        self.executor.submit(detect_job, self.db_path, self.data_dir, job_id)

    def submit_parse(self, job_id: str) -> None:
        self.executor.submit(parse_job, self.db_path, self.data_dir, job_id)
